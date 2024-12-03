# import streamlit as st
# import pandas as pd
# import pdfplumber
# import docx
# import io
# from PIL import Image
# import plotly.express as px

# class DocumentLoader:
#     @staticmethod
#     def load_document(uploaded_file):
#         """
#         Load and preprocess different document types

#         Args:
#             uploaded_file (UploadedFile): Streamlit uploaded file

#         Returns:
#             Preprocessed document content
#         """
#         filename = uploaded_file.name
#         file_extension = filename.split('.')[-1].lower()

#         try:
#             # Text files
#             if file_extension == 'txt':
#                 return uploaded_file.getvalue().decode('utf-8')

#             # PDF files
#             elif file_extension == 'pdf':
#                 text = ""
#                 with pdfplumber.open(io.BytesIO(uploaded_file.getvalue())) as pdf:
#                     for page in pdf.pages:
#                         text += page.extract_text()
#                 return text

#             # Word documents
#             elif file_extension == 'docx':
#                 doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
#                 return '\n'.join([para.text for para in doc.paragraphs])

#             # Excel and CSV files
#             elif file_extension in ['xlsx', 'csv']:
#                 df = pd.read_csv(uploaded_file) if file_extension == 'csv' else pd.read_excel(uploaded_file)
#                 return df

#             # Image files
#             elif file_extension in ['jpg', 'jpeg', 'png', 'gif']:
#                 return Image.open(io.BytesIO(uploaded_file.getvalue()))

#             else:
#                 raise ValueError(f"Unsupported file type: {file_extension}")

#         except Exception as e:
#             st.error(f"Error processing document: {e}")
#             return None

#     @staticmethod
#     def generate_basic_visualization(data):
#         """
#         Generate basic visualization based on data type

#         Args:
#             data (pandas.DataFrame): Input data

#         Returns:
#             Plotly figure or None
#         """
#         if isinstance(data, pd.DataFrame):
#             numeric_columns = data.select_dtypes(include=['int64', 'float64']).columns
#             if len(numeric_columns) > 0:
#                 fig = px.line(data, x=data.index, y=numeric_columns, title='Data Trends')
#                 return fig
#         return None


import streamlit as st
import pandas as pd
import pdfplumber
import docx
import io
from PIL import Image

class DocumentLoader:
    @staticmethod
    def load_document(uploaded_file):
        filename = uploaded_file.name
        file_extension = filename.split('.')[-1].lower()

        try:
            if file_extension == 'txt':
                return uploaded_file.getvalue().decode('utf-8')

            elif file_extension == 'pdf':
                text = ""
                with pdfplumber.open(io.BytesIO(uploaded_file.getvalue())) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text()
                return text

            elif file_extension == 'docx':
                doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
                return '\n'.join([para.text for para in doc.paragraphs])

            elif file_extension in ['xlsx', 'csv']:
                df = pd.read_csv(uploaded_file) if file_extension == 'csv' else pd.read_excel(uploaded_file)
                return df

            elif file_extension in ['jpg', 'jpeg', 'png']:
                return Image.open(io.BytesIO(uploaded_file.getvalue()))

            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

        except Exception as e:
            st.error(f"Error processing document: {e}")
            return None
